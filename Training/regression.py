import os
import time

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import seaborn as sns
from sklearn.model_selection import ParameterGrid, TimeSeriesSplit
from sklearn.preprocessing import StandardScaler
from boruta import BorutaPy
from docx import Document
import statsmodels.api as sm
from statsmodels.stats.outliers_influence import variance_inflation_factor
from sklearn.ensemble import RandomForestRegressor

from models.Cubist import train_and_evaluate as train_cubist
from models.GBM import train_and_evaluate as train_gbm
from models.MARS import train_and_evaluate as train_mars
from models.RF import train_and_evaluate as train_rf
from models.SVM import train_and_evaluate as train_svm
from models.XGB import train_and_evaluate as train_xgb

from models.Cubist import get_parameters as get_cubist_params
from models.MARS import get_parameters as get_mars_params
from models.GBM import get_parameters as get_gbm_params
from models.RF import get_parameters as get_rf_params
from models.SVM import get_parameters as get_svm_params
from models.XGB import get_parameters as get_xgb_params

import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="joblib")
warnings.filterwarnings("ignore", category=RuntimeWarning)


def feature_selection_workflow(X, y, vif_threshold=5.0):
    """Standardize, check VIF, then apply Boruta for feature selection."""
    print("Starting feature selection workflow...")

    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)
    X_scaled_df = pd.DataFrame(X_scaled, columns=X.columns, index=X.index)
    print(f"Features standardized: {X_scaled_df.shape[1]} features")

    filtered_features, removed_by_vif = drop_high_vif_features(
        X_scaled_df, thresh=vif_threshold, protected_features=["tmp"]
    )
    X_filtered = X_scaled_df[filtered_features]
    print(f"---- After VIF filtering: {len(filtered_features)} features remaining")

    model = RandomForestRegressor(n_jobs=-1, random_state=42)
    selector = BorutaPy(model, n_estimators=1000, max_iter=50, perc=90, alpha=0.05, two_step=False, verbose=0)

    y_np = y.to_numpy() if isinstance(y, pd.Series) else np.array(y)
    selector.fit(X_filtered.values, y_np)

    boruta_selected = X_filtered.columns[selector.support_].tolist()
    boruta_rejected = X_filtered.columns[~selector.support_].tolist()
    print(f"---- After Boruta: {len(boruta_selected)} features selected")

    final_features = temperature_rule(boruta_selected)
    print("Selected features:", final_features)

    return {
        'selected_features': final_features,
        'removed_by_vif': removed_by_vif,
        'rejected_by_boruta': boruta_rejected,
        'original_count': X.shape[1],
        'final_count': len(final_features)
    }


def drop_high_vif_features(X, thresh=5.0, protected_features=None):
    """Drop features with high VIF iteratively, preserving protected features."""
    if protected_features is None:
        protected_features = []

    X_copy = X.copy()
    removed_features = []

    while True:
        vif = pd.DataFrame()
        vif["feature"] = X_copy.columns
        vif["VIF"] = [variance_inflation_factor(X_copy.values, i) for i in range(X_copy.shape[1])]

        candidates = vif[~vif["feature"].isin(protected_features)]

        if candidates.empty:
            break

        max_vif = candidates["VIF"].max()
        if max_vif > thresh:
            drop_feature = candidates.loc[candidates["VIF"].idxmax(), "feature"]
            removed_features.append({'feature': drop_feature, 'vif': max_vif, 'reason': 'High VIF'})
            print(f"Dropping '{drop_feature}' (VIF = {max_vif:.2f})")
            X_copy = X_copy.drop(columns=[drop_feature])
        else:
            break

    final_vif = pd.DataFrame()
    final_vif["feature"] = X_copy.columns
    final_vif["VIF"] = [variance_inflation_factor(X_copy.values, i) for i in range(X_copy.shape[1])]
    print("\nVIF scores after dropping high VIF features:")
    print(final_vif)

    if removed_features:
        print(f"\nRemoved {len(removed_features)} features due to high VIF:")
        for item in removed_features:
            print(f"  - {item['feature']}: VIF = {item['vif']:.2f}")

    return X_copy.columns.tolist(), removed_features


def temperature_rule(selected_features):
    """Force tmp to replace individual temperature indicators (tmn, tmx, dtr)."""
    temp_indicators = ['tmn', 'tmx', 'dtr']
    removed = [feature for feature in temp_indicators if feature in selected_features]

    selected_features = [f for f in selected_features if f not in temp_indicators]

    if removed and 'tmp' not in selected_features:
        print("Removed:", removed, "| Added: tmp")
        selected_features.append('tmp')

    return selected_features


def run_ols(X, y, alpha=0.05):
    """Perform OLS regression and print summary and significant predictors."""
    X_ols = sm.add_constant(X)
    model = sm.OLS(y, X_ols).fit()
    print("\n--- OLS Regression Summary ---")
    print(model.summary())

    print(f"\n--- Predictors Significant at alpha = {alpha} ---")
    sig = model.pvalues[model.pvalues < alpha]
    print(sig)

    return model


def calculate_vif(X):
    """Calculate VIF for each feature to detect multicollinearity."""
    vif_data = pd.DataFrame()
    vif_data['Feature'] = X.columns
    vif_data['VIF'] = [variance_inflation_factor(X.values, i) for i in range(X.shape[1])]

    print("\n--- VIF Scores ---")
    print(vif_data)
    return vif_data


def save_results_to_word(results, feature_selection_info=None, filename="model_results_with_fs.docx"):
    """Save training results and feature selection info to a Word document."""
    try:
        results_dir = os.path.join(os.path.dirname(__file__), 'results')
        os.makedirs(results_dir, exist_ok=True)
        file_path = os.path.join(results_dir, filename)
        doc = Document()
        doc.add_heading('Model Training and Performance Results', level=1)

        if feature_selection_info:
            doc.add_heading('Feature Selection Summary per Fold', level=2)

            for fold_idx, fs_info in enumerate(feature_selection_info):
                doc.add_heading(f"Fold {fold_idx + 1}", level=3)

                doc.add_paragraph(f"Original features: {fs_info.get('original_count', 'N/A')}")
                doc.add_paragraph(f"Final selected features: {fs_info.get('final_count', 'N/A')}")

                selected_features = fs_info.get('selected_features', [])
                if selected_features:
                    doc.add_paragraph("Selected Features:", style="List Bullet")
                    for feature in selected_features:
                        doc.add_paragraph(f"  {feature}", style="List Bullet 2")

                removed_by_vif = fs_info.get('removed_by_vif', [])
                if removed_by_vif:
                    doc.add_paragraph("Features Removed by VIF Filtering:", style="List Bullet")
                    for item in removed_by_vif:
                        doc.add_paragraph(f"  {item['feature']} (VIF: {item['vif']:.2f})", style="List Bullet 2")

                rejected_by_boruta = fs_info.get('rejected_by_boruta', [])
                if rejected_by_boruta:
                    doc.add_paragraph("Features Rejected by Boruta:", style="List Bullet")
                    for feature in rejected_by_boruta:
                        doc.add_paragraph(f"  {feature}", style="List Bullet 2")

                doc.add_paragraph()

        for model_type, model_data in results.items():
            doc.add_heading(f"{model_type.upper()} Models", level=2)

            for model_name, metrics_data in model_data.items():
                doc.add_heading(f"{model_name} Performance", level=3)

                doc.add_paragraph(f"Best Average R2: {metrics_data.get('best_avg_r2', 'N/A'):.4f}")

                best_params = metrics_data.get('best_params', {})
                doc.add_paragraph("Best Parameters:", style="List Bullet")
                for param, value in best_params.items():
                    doc.add_paragraph(f"{param}: {value}", style="List Bullet 2")

                if 'avg_metrics' in metrics_data:
                    doc.add_paragraph("Average Metrics:", style="List Bullet")
                    for metric, value in metrics_data['avg_metrics'].items():
                        doc.add_paragraph(f"{metric}: {value:.4f}", style="List Bullet 2")

                if 'fold_r2_scores' in metrics_data:
                    doc.add_paragraph("R2 per Fold:", style="List Bullet")
                    for i, r2 in enumerate(metrics_data['fold_r2_scores'], 1):
                        doc.add_paragraph(f"Fold {i}: {r2:.4f}", style="List Bullet 2")

                doc.add_paragraph()

        doc.save(file_path)
        print(f"\nResults successfully saved to {file_path}")
        return True
    except Exception as e:
        print(f"\nError saving results: {str(e)}")
        return False


def visualize_model_performance(results, with_fs):
    """Create bar chart visualizations for model performance and save to plots/."""
    performance_data = []

    for model_type, model_data in results.items():
        for model_name, metrics in model_data.items():
            avg_metrics = metrics.get('avg_metrics', {})
            performance_data.append({
                'Model Type': model_type,
                'Model': model_name,
                'R2': avg_metrics.get('R2', np.nan),
                'RMSE': avg_metrics.get('RMSE', np.nan),
                'MSE': avg_metrics.get('MSE', np.nan),
                'MAE': avg_metrics.get('MAE', np.nan),
                'MAPE': avg_metrics.get('MAPE', np.nan)
            })

    df_performance = pd.DataFrame(performance_data)

    plot_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'plots')
    os.makedirs(plot_dir, exist_ok=True)

    if 'regression' in results:
        regression_data = df_performance[df_performance['Model Type'] == 'regression']
        suffix = "with_fs" if with_fs else "without_fs"

        for metric in ['R2', 'RMSE', 'MSE', 'MAE']:
            plt.figure(figsize=(12, 6))
            sns.barplot(x='Model', y=metric, data=regression_data)
            plt.title(f'{metric} Comparison - Regression Models', fontsize=14)
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            plt.savefig(os.path.join(plot_dir, f'regression_{metric.lower()}_comparison_{suffix}.png'), dpi=300)
            plt.close()

    # Heatmap of normalized metrics
    heatmap_data = df_performance.set_index('Model').drop(columns='Model Type')
    existing_metrics = [col for col in ['R2', 'RMSE', 'MSE', 'MAE'] if col in heatmap_data.columns]
    heatmap_data = heatmap_data[existing_metrics]

    normalized_data = (heatmap_data - heatmap_data.min()) / (heatmap_data.max() - heatmap_data.min())
    normalized_data.columns.name = "Metrics"

    plt.figure(figsize=(14, 8))
    sns.heatmap(
        normalized_data,
        annot=heatmap_data,
        cmap='YlGnBu',
        fmt='.3f',
        cbar_kws={'label': 'Normalized Value'}
    )
    plt.title('Model Performance Comparison', fontsize=16)
    plt.tight_layout()
    plt.savefig(os.path.join(plot_dir, 'model_performance_heatmap.png'), dpi=300)
    plt.close()

    return df_performance


def load_data(file_name, sheet_name="Sheet1"):
    """Load data from an Excel file in the data/ subdirectory."""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(script_dir, 'data', file_name)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File '{file_name}' not found in {script_dir}")
    return pd.read_excel(file_path, sheet_name=sheet_name)


def check_multicollinearity(X, vif_threshold=10, corr_threshold=0.95):
    """Check multicollinearity via VIF and correlation matrix."""
    vif_df = pd.DataFrame()
    vif_df['Feature'] = X.columns
    vif_df['VIF'] = [variance_inflation_factor(X.values, i) for i in range(X.shape[1])]
    high_vif = vif_df[vif_df['VIF'] > vif_threshold]

    corr_matrix = X.corr().abs()
    upper = corr_matrix.where(np.triu(np.ones(corr_matrix.shape), k=1).astype(bool))
    high_corr = [(col, row, upper.loc[row, col])
                 for col in upper.columns
                 for row in upper.index
                 if upper.loc[row, col] > corr_threshold]

    if not high_vif.empty:
        print(f"\nFeatures with VIF > {vif_threshold}:\n{high_vif}")
    else:
        print(f"\nNo features exceed VIF threshold of {vif_threshold}")

    if high_corr:
        print(f"\nHighly correlated feature pairs (corr > {corr_threshold}):")
        for a, b, val in high_corr:
            print(f"   {a} & {b} -> correlation = {val:.3f}")
    else:
        print(f"\nNo feature pairs exceed correlation threshold of {corr_threshold}")

    return high_vif, high_corr


def main(model_type, fs):
    """Train and evaluate all models with optional feature selection."""
    print(f"\nStarting model training for: {model_type}\n")

    data = load_data('banana_yield_2010-2024.xlsx')
    data['yield'] = data['yield'].replace('#DIV/0!', pd.NA)
    data['yield'] = pd.to_numeric(data['yield'], errors='coerce')
    data = data.dropna(subset=['yield'])
    data = data.sort_values(by='year', ascending=True)

    model_functions = {
        'Cubist': train_cubist,
        'GBM': train_gbm,
        'MARS': train_mars,
        'RF': train_rf,
        'SVM': train_svm,
        'XGB': train_xgb,
    }

    model_params = {
        'Cubist': get_cubist_params,
        'GBM': get_gbm_params,
        'MARS': get_mars_params,
        'RF': get_rf_params,
        'SVM': get_svm_params,
        'XGB': get_xgb_params,
    }

    model_types_to_run = [model_type]
    results = {}

    total_combinations = len(model_types_to_run) * len(model_functions)
    current_combination = 0

    X = data.drop(columns=["yield", "province", "year"])
    y = data["yield"]

    non_zero_features = (X != 0).any(axis=1)
    X = X[non_zero_features]
    y = y[non_zero_features]

    run_ols(X, y)
    calculate_vif(X)

    tscv = TimeSeriesSplit(n_splits=5)
    check_multicollinearity(X)

    if fs == "yes":
        selected_features_per_fold = []
        feature_selection_info_per_fold = []
        for fold_idx, (train_index, test_index) in enumerate(tscv.split(X)):
            print(f"Fold {fold_idx + 1}:")
            X_fold_train = X.iloc[train_index]
            y_fold_train = y.iloc[train_index]

            selected = feature_selection_workflow(X_fold_train, y_fold_train)
            selected_features_per_fold.append(selected['selected_features'])
            feature_selection_info_per_fold.append(selected)

            print(f"  Selected features: {selected}")
            print(f"  Train size: {len(train_index)}")
            print(f"  Test size: {len(test_index)}")
            print("-" * 30)
    elif fs == "no":
        selected_features = list(X.columns)
        print(f"Selected features: {selected_features}")

    for current_model_type in model_types_to_run:
        results[current_model_type] = {}
        for model_name, train_func in model_functions.items():
            current_combination += 1
            start_time = time.time()
            print(f"\nTraining {current_combination} of {total_combinations}: {model_name} for {current_model_type}...")

            best_avg_r2 = -np.inf
            best_params = None
            all_avg_r2_results = []
            param_grid = model_params[model_name]()

            for params in ParameterGrid(param_grid):
                fold_r2_scores = []

                for fold_idx, (train_index, test_index) in enumerate(tscv.split(X)):
                    if fs == "yes":
                        selected_features = selected_features_per_fold[fold_idx]
                    X_fold_train = X.iloc[train_index][selected_features]
                    X_fold_test = X.iloc[test_index][selected_features]
                    y_fold_train = y.iloc[train_index]
                    y_fold_test = y.iloc[test_index]

                    scaler = StandardScaler()
                    X_fold_train_scaled = scaler.fit_transform(X_fold_train)
                    X_fold_test_scaled = scaler.transform(X_fold_test)
                    X_fold_train_scaled = pd.DataFrame(X_fold_train_scaled, columns=selected_features)
                    X_fold_test_scaled = pd.DataFrame(X_fold_test_scaled, columns=selected_features)

                    model, metrics = train_func(
                        X_fold_train_scaled, X_fold_test_scaled, y_fold_train, y_fold_test, params
                    )

                    if model is not None:
                        test_r2 = model.score(X_fold_test_scaled, y_fold_test)
                        fold_r2_scores.append(test_r2)

                if fold_r2_scores:
                    avg_r2 = np.mean(fold_r2_scores)
                    all_avg_r2_results.append((params, avg_r2))

                    if avg_r2 > best_avg_r2:
                        best_avg_r2 = avg_r2
                        best_params = params

            print(f"Best Params for {model_name}: {best_params}")
            print(f"Best Avg R2: {best_avg_r2:.4f}")
            print(f"Execution Time for {model_name}: {time.time() - start_time:.2f} seconds")

            results[current_model_type][model_name] = {
                'best_avg_r2': best_avg_r2,
                'best_params': best_params,
                'all_results': all_avg_r2_results,
                'avg_metrics': {
                    'R2': best_avg_r2,
                    'RMSE': metrics['RMSE'],
                    'MSE': metrics['MSE'],
                    'MAE': metrics['MAE'],
                    'MAPE': metrics['MAPE']
                },
                'fold_r2_scores': fold_r2_scores,
                'fold_best_params': [best_params for _ in range(len(fold_r2_scores))]
            }

    print("\nTraining Complete. Model Performance:")
    for model_type, model_data in results.items():
        print(f"\n--- {model_type.upper()} Models ---")
        for model_name, metrics in model_data.items():
            print(f"\n{model_name} Performance:")
            metrics_to_print = {k: v for k, v in metrics.get('avg_metrics', {}).items() if k != 'residuals'}
            print(f"  Metrics: {metrics_to_print}")

    if fs == "yes":
        return results, data, feature_selection_info_per_fold
    elif fs == "no":
        return results, data


if __name__ == "__main__":
    start_time = time.time()
    fs = "yes"

    if fs == "yes":
        results, data, feature_selection_info_per_fold = main(model_type='regression', fs="yes")
        performance_df = visualize_model_performance(results, True)
        save_results_to_word(results, feature_selection_info_per_fold)
    elif fs == "no":
        results, data = main(model_type='regression', fs="no")
        performance_df = visualize_model_performance(results, False)
        save_results_to_word(results, None, "model_results_without_fs.docx")

    print(f"Overall Execution Time: {time.time() - start_time:.2f} seconds")
