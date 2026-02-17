"""
Export training results to Word document.
"""
import os
from docx import Document
from docx.shared import Pt


def _add_table(doc, df, title=None):
    """Add a pandas DataFrame as a Word table."""
    if title:
        doc.add_paragraph(title, style="List Bullet")

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'

    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)

    for _, row_data in df.iterrows():
        row = table.add_row()
        for i, val in enumerate(row_data):
            if isinstance(val, float):
                row.cells[i].text = f"{val:.4f}"
            else:
                row.cells[i].text = str(val)

    doc.add_paragraph()


def _add_diagnostics(doc, diagnostics):
    """Add initial diagnostics section (OLS, VIF, multicollinearity)."""
    doc.add_heading('Initial Diagnostics (Full Dataset, All 17 Features)', level=2)

    # OLS Summary
    doc.add_heading('OLS Regression Summary', level=3)
    doc.add_paragraph(f"R\u00b2: {diagnostics['ols_r2']:.4f}")
    doc.add_paragraph(f"Adjusted R\u00b2: {diagnostics['ols_adj_r2']:.4f}")

    ols_para = doc.add_paragraph()
    run = ols_para.add_run(diagnostics['ols_summary'])
    run.font.size = Pt(7)
    run.font.name = 'Courier New'

    # Significant predictors
    sig = diagnostics.get('ols_significant', {})
    if sig:
        doc.add_paragraph("Significant Predictors (p < 0.05):", style="List Bullet")
        for name, pval in sig.items():
            doc.add_paragraph(f"{name}: p = {pval:.4f}", style="List Bullet 2")

    # Initial VIF scores (all 17 features)
    doc.add_heading('Initial VIF Scores (All Features)', level=3)
    initial_vif = diagnostics.get('initial_vif')
    if initial_vif is not None:
        _add_table(doc, initial_vif)

    # High VIF features
    high_vif = diagnostics.get('high_vif_features')
    if high_vif is not None and not high_vif.empty:
        doc.add_paragraph("Features with VIF > 10:", style="List Bullet")
        for _, row in high_vif.iterrows():
            doc.add_paragraph(f"{row['Feature']}: VIF = {row['VIF']:.2f}", style="List Bullet 2")

    # High correlation pairs
    high_corr = diagnostics.get('high_corr_pairs', [])
    if high_corr:
        doc.add_heading('Highly Correlated Feature Pairs (r > 0.95)', level=3)
        for a, b, val in high_corr:
            doc.add_paragraph(f"{a} & {b}: r = {val:.3f}", style="List Bullet")

    doc.add_paragraph()


def _add_feature_selection(doc, feature_selection_info):
    """Add per-fold feature selection details."""
    doc.add_heading('Feature Selection Summary per Fold', level=2)

    for fold_idx, fs_info in enumerate(feature_selection_info):
        doc.add_heading(f"Fold {fold_idx + 1}", level=3)

        doc.add_paragraph(f"Original features: {fs_info.get('original_count', 'N/A')}")
        doc.add_paragraph(f"Final selected features: {fs_info.get('final_count', 'N/A')}")

        # VIF before filtering
        vif_before = fs_info.get('vif_before')
        if vif_before is not None:
            _add_table(doc, vif_before, "VIF Scores Before Filtering:")

        # Features removed by VIF
        removed_by_vif = fs_info.get('removed_by_vif', [])
        if removed_by_vif:
            doc.add_paragraph("Features Removed by VIF Filtering:", style="List Bullet")
            for item in removed_by_vif:
                doc.add_paragraph(f"{item['feature']} (VIF: {item['vif']:.2f})", style="List Bullet 2")

        # VIF after filtering
        vif_after = fs_info.get('vif_after')
        if vif_after is not None:
            _add_table(doc, vif_after, "VIF Scores After Filtering:")

        # Boruta results
        boruta_selected = fs_info.get('boruta_selected', [])
        if boruta_selected:
            doc.add_paragraph("Features Confirmed by Boruta:", style="List Bullet")
            for feature in boruta_selected:
                doc.add_paragraph(f"{feature}", style="List Bullet 2")

        rejected_by_boruta = fs_info.get('rejected_by_boruta', [])
        if rejected_by_boruta:
            doc.add_paragraph("Features Rejected by Boruta:", style="List Bullet")
            for feature in rejected_by_boruta:
                doc.add_paragraph(f"{feature}", style="List Bullet 2")

        # Final selected features
        selected_features = fs_info.get('selected_features', [])
        if selected_features:
            doc.add_paragraph("Final Selected Features (after temperature rule):", style="List Bullet")
            for feature in selected_features:
                doc.add_paragraph(f"{feature}", style="List Bullet 2")

        doc.add_paragraph()


def save_results_to_word(results, feature_selection_info=None, filename="model_results_with_fs.docx", diagnostics=None):
    """Save training results, diagnostics, and feature selection info to a Word document."""
    try:
        results_dir = os.path.join(os.path.dirname(__file__), 'results')
        os.makedirs(results_dir, exist_ok=True)
        file_path = os.path.join(results_dir, filename)
        doc = Document()
        doc.add_heading('Model Training and Performance Results', level=1)

        # Diagnostics section
        if diagnostics:
            _add_diagnostics(doc, diagnostics)

        # Feature selection section
        if feature_selection_info:
            _add_feature_selection(doc, feature_selection_info)

        # Model results section
        for model_type, model_data in results.items():
            doc.add_heading(f"{model_type.upper()} Models", level=2)

            for model_name, metrics_data in model_data.items():
                doc.add_heading(f"{model_name} Performance", level=3)

                doc.add_paragraph(f"Best Average R2: {metrics_data.get('best_avg_r2', 'N/A'):.4f}")

                best_params = metrics_data.get('best_params', {})
                doc.add_paragraph("Best Parameters:", style="List Bullet")
                for param, value in best_params.items():
                    doc.add_paragraph(f"{param}: {value}", style="List Bullet 2")

                if 'avg_metrics' in metrics_data:
                    doc.add_paragraph("Average Metrics:", style="List Bullet")
                    for metric, value in metrics_data['avg_metrics'].items():
                        doc.add_paragraph(f"{metric}: {value:.4f}", style="List Bullet 2")

                if 'fold_r2_scores' in metrics_data:
                    doc.add_paragraph("R2 per Fold:", style="List Bullet")
                    for i, r2 in enumerate(metrics_data['fold_r2_scores'], 1):
                        doc.add_paragraph(f"Fold {i}: {r2:.4f}", style="List Bullet 2")

                doc.add_paragraph()

        doc.save(file_path)
        print(f"\nResults successfully saved to {file_path}")
        return True
    except Exception as e:
        print(f"\nError saving results: {str(e)}")
        return False
